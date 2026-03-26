import sys
import json
import os
import google.generativeai as genai
from docx import Document

if len(sys.argv) < 3:
    print("Usage: python update_docx.py <path_to_docx> <path_to_ai_resume_txt>")
    sys.exit(1)

docx_path = sys.argv[1]
ai_resume_path = sys.argv[2]
outname = docx_path.replace(".docx", f"_{os.path.basename(ai_resume_path)}.docx") 
# e.g., generalcv_ai_cv_123.txt.docx -> handled later

genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-2.5-flash")

doc = Document(docx_path)
paras = []

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                cleaned = p.text.strip()
                if len(cleaned) > 10 and cleaned not in paras:
                    paras.append(cleaned)
for p in doc.paragraphs:
    cleaned = p.text.strip()
    if len(cleaned) > 10 and cleaned not in paras:
        paras.append(cleaned)

with open(ai_resume_path, "r", encoding="utf-8") as f:
    ai_resume = f.read()

prompt = f"""You are an ATS docx updater. 
I am going to give you a list of original text paragraphs from a .docx file, and a newly optimized AI ATS Resume text.
Find which original paragraphs should be modified to reflect the newly optimized resume and return a strict JSON dictionary mapping.
Format: {{"Original exact paragraph string": "New optimized paragraph string"}}
Only include paragraphs that need changes (skills, tools, project bullets, summary).
Do not change styling or add markdown.

Original Paragraphs extract:
{json.dumps(paras, indent=2)}

New ATS Optimized Resume constraints/text:
{ai_resume}

Return ONLY raw JSON dictionary.
"""

response = model.generate_content(prompt)
out_text = response.text.strip()
if out_text.startswith("```json"):
    out_text = out_text.replace("```json", "", 1)
if out_text.startswith("```"):
    out_text = out_text.replace("```", "", 1)
if out_text.endswith("```"):
    out_text = out_text[:-3]

import re
try:
    mappings = json.loads(out_text.strip())
except json.JSONDecodeError:
    match = re.search(r'\{.*\}', out_text, re.DOTALL)
    if match:
        mappings = json.loads(match.group(0))
    else:
        mappings = {}

# Apply
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                cleaned = p.text.strip()
                if cleaned in mappings:
                    p.text = mappings[cleaned]

for p in doc.paragraphs:
    cleaned = p.text.strip()
    if cleaned in mappings:
        p.text = mappings[cleaned]

# Safe output name processing
folder = os.path.dirname(docx_path)
basename = os.path.basename(docx_path).replace(".docx", "")
ts = os.path.basename(ai_resume_path).replace(".txt", "")
final_out = os.path.join(folder, f"{basename}_{ts}.docx")

doc.save(final_out)
print(final_out)
